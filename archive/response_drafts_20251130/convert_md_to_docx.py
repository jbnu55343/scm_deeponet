import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_markdown_paragraph(doc, text):
    """
    A simple markdown parser for paragraphs.
    Handles bold (**text**) and basic text.
    """
    p = doc.add_paragraph()
    
    # Split by bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)

def convert_md_to_docx(source_md_file, output_docx_file):
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    with open(source_md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Heading 1
            h = doc.add_heading(line[2:], level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = h.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('## '):
            # Heading 2
            h = doc.add_heading(line[3:], level=2)
            run = h.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('### '):
            # Heading 3
            h = doc.add_heading(line[4:], level=3)
            run = h.runs[0]
            run.font.color.rgb = RGBColor(0, 0, 0)
        elif line.startswith('---'):
            doc.add_page_break()
        else:
            # Standard paragraph with simple bold parsing
            add_markdown_paragraph(doc, line)

    doc.save(output_docx_file)
    print(f"Successfully created {output_docx_file}")

# List of files to convert
files = [
    "Response_to_Reviewer_1.md",
    "Response_to_Reviewer_2.md",
    "Response_to_Reviewer_3.md",
    "Response_to_Reviewer_4.md",
    "Response_to_Reviewer_5.md"
]

current_dir = os.path.dirname(os.path.abspath(__file__))

for filename in files:
    md_path = os.path.join(current_dir, filename)
    docx_path = os.path.join(current_dir, filename.replace(".md", ".docx"))
    
    if os.path.exists(md_path):
        try:
            convert_md_to_docx(md_path, docx_path)
        except Exception as e:
            print(f"Error converting {filename}: {e}")
    else:
        print(f"File not found: {md_path}")
